"""
Apply surgical text replacements to a .docx while preserving formatting.

Part of PhD Thesis: Triple-Hybrid RAG with PPO-based L-DWA
Author: Shin Dong-wook <sdw@hoseo.ac.kr>

Each replacement is an exact-string find-and-replace applied at the run
level. If the target text spans multiple runs (mid-sentence formatting),
the replacement is skipped and reported — the user must fix those in Word.

SAFETY: always saves to a new file (_v5.docx by default). Never overwrites.

Usage:
    python scripts/patch_docx.py --input thesis_current/박사논문_1장_서론_v4.docx \
        --output thesis_current/박사논문_1장_서론_v5.docx --patch ch1
"""

from __future__ import annotations

import argparse
import logging
import sys
from pathlib import Path

from docx import Document

logging.basicConfig(format="%(asctime)s [%(levelname)s] %(message)s", level=logging.INFO)
logger = logging.getLogger(__name__)


# ---------- patch rules (exact-string find-and-replace) ----------

PATCHES: dict[str, list[tuple[str, str]]] = {
    # ------------- Ch.1 ----------------
    "ch1": [
        ("Ⅵ장 3·7절 (F1 +3.5%, Boundary EM +32.8%)",
         "Ⅵ장 3·4절 (F1_strict +18.6%, Conditional F1_substring +131%)"),
        ("Ⅴ장 4절 + Ⅵ장 7절 (Boundary EM 0.61→0.81)",
         "Ⅴ장 4절 + Ⅵ장 4절 (Conditional F1_substring 0.090→0.208, +131%)"),
        ("합성 대학 F1 0.89±0.01 (R-DWA 0.86 대비 +3.5%)",
         "합성 대학 F1_strict 0.085±0.001, F1_substring 0.482±0.006 "
         "(R-DWA F1_strict 0.072 대비 +18.6%, F1_substring 0.448 대비 +7.5%, 3 seeds 평균)"),
        ("Ablation (A)→(D) 단조 증가, L-DWA 추가 F1 +3.9%",
         "Ablation (A)→(D) 단조 증가, L-DWA 추가 F1_strict +18.6% (3 seeds 재현성 std < 1.5%)"),
        ("Boundary EM 0.61→0.81 (+32.8% 개선)",
         "Conditional F1_substring 0.090→0.208 (+131% 개선, Oracle 상한 0.195 초과)"),
        ("경계선상 질의 EM +33% 개선",
         "조건부 질의 F1_substring +131% 개선"),
    ],
    # ------------- Ch.5 ----------------
    "ch5": [
        ("0.88 → 0.89±0.01 (최종)",
         "R_early=0.197 → R_late=0.215±0.002 (ep 10,000, 3 seeds)"),
        ("BERT 멀티 레이블은 규칙 기반 대비 정확도 +3.5%p",
         "BERT 멀티 레이블은 규칙 기반 대비 정확도 +3.5%p (본 M6 버전은 intent_logits=0으로 학습하여 density + source_stats + query_meta만으로도 조건부 +131%를 달성)"),
        ("Boundary EM 0.61 → 0.81",
         "조건부 F1_substring 0.090 → 0.208 (+131%)"),
    ],
    # ------------- Ch.7 ----------------
    "ch7": [
        ("F1 +3.5%, Boundary EM +32.8% 추가 개선의 정량적 규명",
         "F1_strict +18.6%, 조건부 F1_substring +131% 개선의 정량적 규명 (3 seeds 재현성 std < 1.5%)"),
        ("경계선상 질의 EM +33%의 개선",
         "조건부 F1_substring +131% 개선"),
        ("EM 0.61 (R-DWA)", "F1_substring 0.090 (R-DWA)"),
        ("EM 0.81 (L-DWA)", "F1_substring 0.208 (L-DWA, Oracle 0.195 초과)"),
    ],
}


# ---------- patching ----------

def patch_document(doc: Document, patches: list[tuple[str, str]]) -> tuple[int, list[str]]:
    """Apply patches, return (n_applied, list_of_skipped_targets)."""
    applied = 0
    skipped: list[str] = []

    for old, new in patches:
        found = False
        for p in doc.paragraphs:
            # Check if any single run contains the exact old text
            for run in p.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    applied += 1
                    found = True
                    break
            if found:
                break
            # Try concatenated-paragraph-text match (spans multiple runs)
            full = "".join(r.text for r in p.runs)
            if old in full:
                # Replace by rewriting the paragraph text into first run and
                # clearing the rest — this WILL lose per-word formatting, so
                # we skip and report instead.
                skipped.append(old)
                found = True
                break
        if not found:
            skipped.append(old)

    # Also check table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        for old, new in patches:
                            if old in run.text:
                                run.text = run.text.replace(old, new)
                                applied += 1

    return applied, skipped


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__,
                                     formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--patch", required=True, choices=list(PATCHES))
    args = parser.parse_args(argv)

    if not args.input.exists():
        logger.error("Input not found: %s", args.input)
        return 1

    doc = Document(str(args.input))
    n_applied, skipped = patch_document(doc, PATCHES[args.patch])

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.output))
    logger.info("Patched %d run-level matches → %s", n_applied, args.output)
    if skipped:
        logger.warning("%d patches SKIPPED (manual fix in Word required):", len(skipped))
        for s in skipped:
            logger.warning("  · %s", s[:100])
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

"""
Convert a thesis-chapter markdown file to a Hoseo-style .docx.

Part of PhD Thesis: Triple-Hybrid RAG with PPO-based L-DWA
Author: Shin Dong-wook <sdw@hoseo.ac.kr>

Intended for the Ch.6 v4 markdown draft (``thesis_current/박사논문_6장_실험평가_v4_draft.md``).
Preserves heading hierarchy (# → 장, ## → 절, ### → 하위절), GFM tables,
numbered / bulleted lists, and paragraph text.

Hoseo formatting applied (per 호서대 박사 논문 규격):
- Paper: A5 4.6배판 approximation (188 × 257 mm)
- Font: 한글 바탕체 10.5pt for body, 12pt bold for 절, 14pt bold for 장
- Line spacing: 200% (double) for body
- Margins: top/bottom 25mm, left/right 20mm

Usage:
    python scripts/md_to_thesis_docx.py \\
        --input thesis_current/박사논문_6장_실험평가_v4_draft.md \\
        --output thesis_current/박사논문_6장_실험평가_v4.docx \\
        --chapter "제Ⅵ장 실험 및 평가"

The markdown parser is intentionally minimal — it handles the patterns we
actually use in thesis drafts, not arbitrary GFM.
"""

from __future__ import annotations

import argparse
import logging
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt

logging.basicConfig(format="%(asctime)s [%(levelname)s] %(message)s", level=logging.INFO)
logger = logging.getLogger(__name__)


# ---------- styling ----------

BODY_FONT = "바탕"  # Korean body — Bronze equivalent on most systems
BODY_SIZE = Pt(10.5)
H1_SIZE = Pt(14)
H2_SIZE = Pt(12)
H3_SIZE = Pt(11)
CODE_FONT = "D2Coding"  # fallback handled by Word

PAGE_WIDTH_MM = 188  # 4.6배판 approximation (close to B5-narrow)
PAGE_HEIGHT_MM = 257
MARGIN_TB_MM = 25
MARGIN_LR_MM = 20


def _set_font(run, *, name: str = BODY_FONT, size=BODY_SIZE, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = size
    run.bold = bold
    # Ensure the Korean font is used for East Asian characters too
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def _set_paragraph_spacing(para, *, line: float = 2.0, space_after: float = 0.0) -> None:
    fmt = para.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line
    fmt.space_after = Pt(space_after)


def _configure_page(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Mm(PAGE_WIDTH_MM)
        section.page_height = Mm(PAGE_HEIGHT_MM)
        section.top_margin = Mm(MARGIN_TB_MM)
        section.bottom_margin = Mm(MARGIN_TB_MM)
        section.left_margin = Mm(MARGIN_LR_MM)
        section.right_margin = Mm(MARGIN_LR_MM)


# ---------- markdown parsing ----------

HEADING_RE = re.compile(r"^(#{1,4})\s+(.*?)\s*$")
TABLE_ROW_RE = re.compile(r"^\|(.+)\|\s*$")
TABLE_SEP_RE = re.compile(r"^\|[\s:|-]+\|\s*$")
BULLET_RE = re.compile(r"^(\s*)[-*]\s+(.*)$")
NUMBERED_RE = re.compile(r"^(\s*)(\d+)\.\s+(.*)$")
CODE_FENCE_RE = re.compile(r"^```")
BLOCKQUOTE_RE = re.compile(r"^>\s?(.*)$")
HR_RE = re.compile(r"^-{3,}\s*$")
IMAGE_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")


def _strip_md_inline(text: str) -> str:
    """Remove simple markdown inline: **bold** → bold, `code` → code, [text](url) → text.

    Also strips HTML inline tags commonly used for sub/sup in markdown
    (Word does not render these as tags). Subscript/superscript glyphs are
    folded into the surrounding text inline. This keeps the paragraph readable
    after Word/한글 import, even though true subscript styling is lost; users
    can re-apply Alt+Shift+S after paste if needed.
    """
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)  # links
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)  # **bold**
    text = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", text)  # *italic*
    text = re.sub(r"`([^`]+)`", r"\1", text)  # `code`
    # HTML inline tags — fold contents inline (Word does not render <sub> etc.).
    # Apply repeatedly to handle nested cases like <sub>A<sub>grid</sub></sub>.
    for _ in range(5):  # bounded fixed-point
        new_text = re.sub(r"<sub>([^<]+)</sub>", r"_\1", text)
        new_text = re.sub(r"<sup>([^<]+)</sup>", r"^\1", new_text)
        if new_text == text:
            break
        text = new_text
    text = re.sub(r"</?(b|i|u|em|strong|span|br)\s*/?>", "", text, flags=re.IGNORECASE)
    return text


def _split_table_row(line: str) -> list[str]:
    cells = [c.strip() for c in line.strip("|").split("|")]
    return cells


# ---------- rendering ----------

def _add_heading(doc: Document, text: str, level: int) -> None:
    style_name = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}.get(
        level, "Heading 2"
    )
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(_strip_md_inline(text))
    size = {1: H1_SIZE, 2: H2_SIZE, 3: H3_SIZE}.get(level, H2_SIZE)
    _set_font(run, size=size, bold=True)
    if level == 1:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _set_paragraph_spacing(p, line=1.5, space_after=8)


def _add_paragraph(doc: Document, text: str) -> None:
    if not text.strip():
        return
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, line=2.0, space_after=0)
    text = _strip_md_inline(text)
    run = p.add_run(text)
    _set_font(run)


def _add_bullet(doc: Document, text: str, indent_level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    _set_paragraph_spacing(p, line=1.5, space_after=0)
    run = p.add_run(_strip_md_inline(text))
    _set_font(run)
    if indent_level > 0:
        p.paragraph_format.left_indent = Cm(0.75 * indent_level)


def _add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    _set_paragraph_spacing(p, line=1.5, space_after=0)
    run = p.add_run(_strip_md_inline(text))
    _set_font(run)


def _add_table(doc: Document, rows: list[list[str]], indent_cm: float = 0.0) -> None:
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    if indent_cm > 0:
        # Indent table by setting indent on first paragraph in each cell
        # (python-docx does not expose table-level indent directly; the table
        # itself will appear left-aligned. We use cell indent as a proxy.)
        try:
            from docx.oxml import OxmlElement
            tblPr = table._tbl.tblPr
            tblInd = OxmlElement("w:tblInd")
            tblInd.set(qn("w:w"), str(int(indent_cm * 567)))  # twips ≈ cm × 567
            tblInd.set(qn("w:type"), "dxa")
            tblPr.append(tblInd)
        except Exception:
            pass
    for i, row in enumerate(rows):
        for j in range(n_cols):
            cell = table.cell(i, j)
            cell.text = ""  # clear default
            p = cell.paragraphs[0]
            text = row[j] if j < len(row) else ""
            run = p.add_run(_strip_md_inline(text))
            _set_font(run, size=Pt(9.5), bold=(i == 0))
            _set_paragraph_spacing(p, line=1.15, space_after=0)


def _add_box_paragraph(doc: Document, text: str, indent_cm: float = 1.0) -> None:
    """Indented paragraph (used inside blockquote / box rendering). No ▸ prefix."""
    if not text.strip():
        return
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, line=1.5, space_after=0)
    p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(_strip_md_inline(text))
    _set_font(run, size=Pt(10))


def _add_box_bullet(doc: Document, text: str, indent_cm: float = 1.0) -> None:
    """Indented bullet line for inside a blockquote box."""
    p = doc.add_paragraph(style="List Bullet")
    _set_paragraph_spacing(p, line=1.5, space_after=0)
    p.paragraph_format.left_indent = Cm(indent_cm + 0.5)
    run = p.add_run(_strip_md_inline(text))
    _set_font(run, size=Pt(10))


def _process_blockquote_block(doc: Document, content_lines: list[str]) -> None:
    """Process the contents of a blockquote (markdown box) with full markdown
    parsing — tables, bullets, paragraphs — all rendered as proper Word
    elements with left indent so the box visual is preserved.
    """
    indent_cm = 1.0
    i = 0
    while i < len(content_lines):
        line = content_lines[i]

        # blank
        if not line.strip():
            i += 1
            continue

        # heading inside box (e.g. "**[박스 1-1] ...**" rendered as bold paragraph)
        # we do not promote to true heading — keep as bold indented line
        # since the box is itself visually demarcated.

        # table — accumulate contiguous rows
        if TABLE_ROW_RE.match(line):
            rows_raw = []
            while i < len(content_lines) and TABLE_ROW_RE.match(content_lines[i]):
                if not TABLE_SEP_RE.match(content_lines[i]):
                    rows_raw.append(_split_table_row(content_lines[i]))
                i += 1
            _add_table(doc, rows_raw, indent_cm=indent_cm)
            continue

        # bullet
        mb = BULLET_RE.match(line)
        if mb:
            _add_box_bullet(doc, mb.group(2), indent_cm=indent_cm)
            i += 1
            continue

        # numbered list
        mn = NUMBERED_RE.match(line)
        if mn:
            _add_box_paragraph(doc, f"{mn.group(2)}. {mn.group(3)}", indent_cm=indent_cm)
            i += 1
            continue

        # plain paragraph
        _add_box_paragraph(doc, line, indent_cm=indent_cm)
        i += 1


def _add_blockquote(doc: Document, text: str) -> None:
    """Single-line blockquote helper (legacy). For multi-line blockquote blocks,
    use _process_blockquote_block which handles tables / bullets correctly.
    """
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, line=1.5, space_after=0)
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(_strip_md_inline(text))
    _set_font(run, size=Pt(10))


def _add_image(doc: Document, src: str, caption: str, base_dir: Path) -> None:
    """Embed image from `src` (absolute or relative to md file's dir)."""
    img_path = Path(src)
    if not img_path.is_absolute():
        # Try relative to md dir, then repo root
        candidates = [
            base_dir / src,
            Path(__file__).resolve().parent.parent / src,
        ]
        img_path = next((p for p in candidates if p.exists()), candidates[0])
    if not img_path.exists():
        logger.warning("image not found: %s (skipped)", img_path)
        return
    try:
        # Fit to page width minus small margin (≈14cm of 14.8cm text area)
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Cm(14.0))
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            _set_paragraph_spacing(cap, line=1.15, space_after=6)
            cap_run = cap.add_run(caption)
            _set_font(cap_run, size=Pt(9.5), bold=True)
    except Exception as e:  # noqa: BLE001
        logger.warning("failed to embed %s: %s", img_path, e)


def convert_markdown_to_docx(md_path: Path, out_path: Path, chapter_title: str | None = None) -> None:
    doc = Document()
    _configure_page(doc)

    if chapter_title:
        _add_heading(doc, chapter_title, level=1)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        # code fence
        if CODE_FENCE_RE.match(line):
            if in_code:
                # flush code block
                p = doc.add_paragraph()
                _set_paragraph_spacing(p, line=1.15, space_after=0)
                run = p.add_run("\n".join(code_buf))
                _set_font(run, name=CODE_FONT, size=Pt(9))
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # heading
        m = HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            _add_heading(doc, m.group(2), level)
            i += 1
            continue

        # blockquote — accumulate ALL contiguous '> ...' lines into a buffer,
        # then process the buffer with full markdown parsing (tables, bullets,
        # paragraphs) so that box contents render correctly.
        if BLOCKQUOTE_RE.match(line):
            bq_buf: list[str] = []
            while i < len(lines):
                bm = BLOCKQUOTE_RE.match(lines[i])
                if bm:
                    bq_buf.append(bm.group(1))
                    i += 1
                    continue
                # Allow a single blank line between '>' blocks to be part of
                # the same box (markdown convention).
                if not lines[i].strip() and i + 1 < len(lines) and BLOCKQUOTE_RE.match(lines[i + 1]):
                    bq_buf.append("")
                    i += 1
                    continue
                break
            _process_blockquote_block(doc, bq_buf)
            continue

        # image — "![alt](path)" on its own line
        img_m = IMAGE_RE.match(line)
        if img_m:
            alt, path_str = img_m.group(1), img_m.group(2)
            _add_image(doc, path_str, alt, md_path.parent)
            i += 1
            continue

        # hr / blank
        if HR_RE.match(line) or not line.strip():
            i += 1
            continue

        # table — accumulate contiguous table rows
        if TABLE_ROW_RE.match(line):
            rows_raw = []
            while i < len(lines) and TABLE_ROW_RE.match(lines[i]):
                if not TABLE_SEP_RE.match(lines[i]):
                    rows_raw.append(_split_table_row(lines[i]))
                i += 1
            _add_table(doc, rows_raw)
            continue

        # bullet
        mb = BULLET_RE.match(line)
        if mb:
            indent = len(mb.group(1)) // 2
            _add_bullet(doc, mb.group(2), indent_level=indent)
            i += 1
            continue

        # numbered
        mn = NUMBERED_RE.match(line)
        if mn:
            _add_numbered(doc, mn.group(3))
            i += 1
            continue

        # plain paragraph
        _add_paragraph(doc, line)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    logger.info("Wrote %s", out_path)


# ---------- CLI ----------

def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__,
                                     formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--chapter", default=None,
                        help="Optional chapter title to prepend (centered, level-1).")
    args = parser.parse_args(argv)

    convert_markdown_to_docx(args.input, args.output, args.chapter)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
